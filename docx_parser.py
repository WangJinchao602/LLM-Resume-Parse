'''
这是读取.docx文件
Author: wangjinchao
date: 2025/08/07
'''
import os

from bs4 import BeautifulSoup


# 普通文本、文本框、表格分开提取
# 由于简历大部分会使用模板，所有基本都在文本框内，或则表格内
# 时间快
def extract_text_split(docx_path):
    from docx import Document
    from lxml import etree
    import zipfile

    # 加载.docx文件
    source_docx = docx_path
    document = Document(source_docx)

    # 解压.docx文件以访问其XML内容
    with zipfile.ZipFile(source_docx, 'r') as docx_zip:
        xml_content = docx_zip.read('word/document.xml')

    # 解析XML内容
    xml_tree = etree.XML(xml_content)

    # 定义命名空间
    namespaces = {
        'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main',
        'v': 'urn:schemas-microsoft-com:vml'
    }

    # 已提取段落的文本集合
    extracted_texts = []

    # 1. 提取普通段落
    for paragraph in document.paragraphs:
        para_text = paragraph.text
        if para_text:
            # print(para_text)
            extracted_texts.append(para_text)

    # 2. 提取文本框中的段落
    for textbox in xml_tree.findall('.//w:txbxContent', namespaces):
        for paragraph in textbox.findall('.//w:p', namespaces):
            # 提取段落中的所有文本
            texts = []
            for text_elem in paragraph.findall('.//w:t', namespaces):
                if text_elem.text:
                    texts.append(text_elem.text)

            para_text = ''.join(texts).strip()

            # 检查段落文本是否唯一且非空
            if para_text and para_text not in extracted_texts:
                # print(para_text)
                extracted_texts.append(para_text)

    # 提取表格中的文本
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                extracted_texts.append(cell.text)

    return '\n'.join(extracted_texts)

from ordered_set import OrderedSet
def get_paragraphs_text_doc(path):
    # 打开docx文件
    document = zipfile.ZipFile(path)
    xml = document.read("word/document.xml")
    wordObj = BeautifulSoup(xml.decode("utf-8"), "lxml")
    # 找到所有的<w:t>、<w:br/>和<w:p>标签，这些标签包含了文字内容、换行符和段落符号
    text_tags = wordObj.find_all(["w:t", "w:p"])
    # 保存提取的文字
    extracted_text = ""
    # 遍历每个标签
    for tag in text_tags:
        # 如果是<w:t>标签，提取文字内容
        if tag.name == "w:t":
            text = tag.text.strip()
            text = text.replace(" ", "")
            if text:
                extracted_text += text
        # 如果是<w:p>标签，添加段落符号（仅在非空行时添加）
        elif tag.name == "w:p":
            if extracted_text and extracted_text[-1] != "\n":
                extracted_text += "\n"

    lines = extracted_text.split("\n")
    my_ordered_set = OrderedSet(lines)
    # 输出提取的文字
    newtext = "\n".join(my_ordered_set)
    return newtext

# 时间慢
"""
    从Word文档中提取所有文本内容（统一XML解析方案）
    包括：普通段落、文本框、表格、页眉、页脚、脚注、尾注等
"""
from lxml import etree
import zipfile
import re
from collections import OrderedDict
def extract_text_from_docx(docx_path):
    # 存储所有唯一文本内容
    unique_texts = OrderedDict()

    # 定义Word文档命名空间
    namespaces = {
        'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main',
        'v': 'urn:schemas-microsoft-com:vml',
        'wp': 'http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing',
        'mc': 'http://schemas.openxmlformats.org/markup-compatibility/2006'
    }

    # 需要提取文本的XML文件列表
    xml_files = [
        'word/document.xml',  # 主文档内容
        'word/header[0-9]*.xml',  # 所有页眉
        'word/footer[0-9]*.xml',  # 所有页脚
        'word/footnotes.xml',  # 脚注
        'word/endnotes.xml',  # 尾注
    ]

    try:
        with zipfile.ZipFile(docx_path, 'r') as docx_zip:
            # 获取文档中所有文件列表
            file_list = docx_zip.namelist()

            # 处理所有相关的XML文件
            for pattern in xml_files:
                # 匹配符合模式的文件
                matched_files = [f for f in file_list if re.match(pattern, f)]

                for xml_file in matched_files:
                    try:
                        # 读取并解析XML内容
                        xml_content = docx_zip.read(xml_file)
                        xml_tree = etree.XML(xml_content)

                        # 提取该XML文件中的所有文本
                        extract_text_from_xml(xml_tree, namespaces, unique_texts)

                    except Exception as e:
                        print(f"处理文件 {xml_file} 时出错: {str(e)}")
                        continue

            # 特殊处理：文本框内容可能在vmlDrawing文件中
            vml_files = [f for f in file_list if f.startswith('word/drawing') and f.endswith('.xml')]
            for vml_file in vml_files:
                try:
                    xml_content = docx_zip.read(vml_file)
                    xml_tree = etree.XML(xml_content)

                    # 提取VML绘图中的文本框内容
                    extract_textboxes_from_vml(xml_tree, namespaces, unique_texts)

                except Exception as e:
                    print(f"处理VML文件 {vml_file} 时出错: {str(e)}")
                    continue

    except Exception as e:
        print(f"打开或读取文档时出错: {str(e)}")
        return []

    # 返回提取的所有唯一文本（保持顺序）
    return list(unique_texts.keys())


def extract_text_from_xml(xml_tree, namespaces, unique_texts):
    """
    从XML树中提取所有文本内容
    """
    # 提取所有段落（包括普通段落和表格中的段落）
    paragraphs = xml_tree.xpath('//w:p', namespaces=namespaces)

    for para in paragraphs:
        # 提取段落中的所有文本节点
        text_elements = para.xpath('.//w:t', namespaces=namespaces)

        # 拼接段落文本
        para_text = ''.join([elem.text if elem.text else '' for elem in text_elements]).strip()

        # 移除无意义的控制字符
        para_text = re.sub(r'[\x00-\x1F\x7F]', '', para_text)

        # 保存非空唯一文本
        if para_text and para_text not in unique_texts:
            unique_texts[para_text] = True

    # 提取文本框内容（如果存在）
    textboxes = xml_tree.xpath('//w:txbxContent', namespaces=namespaces)
    for textbox in textboxes:
        extract_text_from_xml(textbox, namespaces, unique_texts)

    # 提取表格中的文本
    tables = xml_tree.xpath('//w:tbl', namespaces=namespaces)
    for table in tables:
        extract_text_from_xml(table, namespaces, unique_texts)


def extract_textboxes_from_vml(xml_tree, namespaces, unique_texts):
    """
    从VML绘图文件中提取文本框内容
    """
    # 查找所有文本框
    textboxes = xml_tree.xpath('//v:textbox', namespaces=namespaces)

    for textbox in textboxes:
        # 获取文本框内容
        content = textbox.xpath('.//w:txbxContent', namespaces=namespaces)
        if content:
            extract_text_from_xml(content[0], namespaces, unique_texts)


# 3. docx2pdf
# def docx2pdf_ocr(docx_path):
#     # 转换为pdf
#     from docx2pdf import convert
#     cache_folder = 'cache'
#     if not os.path.exists(cache_folder):
#         os.makedirs(cache_folder)
#     convert(docx_path, cache_folder+'\\cache.pdf')
#
#     # 解析pdf
#     from pdf_parser import spire_pdf_ocr
#     text = spire_pdf_ocr(cache_folder+'\\cache.pdf')
#     # 删除缓存文件
#     os.remove(cache_folder+'\\cache.pdf')
#     return text

# 4. docx2png
def docx2png_ocr(docx_path):
    cache_folder = 'cache'
    # 转换为pdf
    import aspose.words as aw
    doc = aw.Document(docx_path)
    options = aw.saving.ImageSaveOptions(aw.SaveFormat.PNG)
    for pageNumber in range(doc.page_count):
        options.page_set = aw.saving.PageSet(pageNumber)
        png_path = os.path.join(cache_folder, f"{pageNumber + 1}_page.png")
        doc.save(png_path, options)

    # 解析pdf
    # 遍历读取png图片并进行OCR识别
    all_text = []
    # 按页码顺序处理图片
    for pageNumber in range(1, doc.page_count + 1):
        png_path = os.path.join(cache_folder, f"{pageNumber}_page.png")
        if os.path.exists(png_path):
            # 使用OCR识别图片中的文字
            try:
                # OCR
                from image_ocr import cn_ocr
                page_text = cn_ocr(png_path)
                all_text.append(page_text)

            except Exception as e:
                print(f"OCR识别第{pageNumber}页失败: {e}")
                all_text.append("")  # 添加空字符串保持页码顺序

    # 合并所有页面的文本
    final_text = '\n'.join(all_text)

    # 删除缓存文件
    for pageNumber in range(1, doc.page_count + 1):
        png_path = os.path.join(cache_folder, f"{pageNumber}_page.png")
        if os.path.exists(png_path):
            os.remove(png_path)

    return final_text


if __name__ == "__main__":
    import time
    from request_llm import request_llm

    image_path = 'data\\docx'
    # 遍历文件夹下的所有文件
    for file in os.listdir(image_path):
        # 拼接文件路径
        file_path = os.path.join(image_path, file)
        print(file_path)
        # 直接提取
        stat_time = time.time()
        text = get_paragraphs_text_doc(file_path)
        end_time = time.time()
        print(text)
        # 调用LLM
        prompt = f'{text}'
        # result = request_llm(prompt)
        # 将识别结果存到文件,追加
        with open('data\\docx\\docx-result.txt', 'a+', encoding='utf-8') as f:
            f.write(text)
            f.write('\n')
            # f.write(result)
            # f.write('\n')
            f.write('=' * 20)
        f.close()

